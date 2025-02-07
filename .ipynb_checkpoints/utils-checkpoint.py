# -*- coding: utf-8 -*-

import numpy as np
import pandas as pd
from docx import Document
from SPARQLWrapper import SPARQLWrapper, JSON
from tqdm.notebook import tqdm

######################################################

# WRITE, SEND AND AGGREGATE THE QUERIES

def parse_list(iiif_list):
    """
    This function takes a list of URLs/URIs as a TXT file
    with one URL/URI per line. Anything which does not
    contain bnf.fr will simply be ignored.

    The output is a human-readable XLSX file with all information sorted.

    :param iiif_list: The path to the TXT file containing the URL/URI list.
    
    """

    # Initiate the base list.
    all_dfs = []

    # Get the URI/URL list from the file.
    with open(iiif_list) as f:
        texte = f.read()
    
    links = texte.split("\n")

    # For each line in the file.
    for l in tqdm(links):

        # If it is a Gallica URL:
        
        if "http://gallica.bnf.fr" in l:

            # Define the main parts of the query.
            
            # Beginning of the query.
            qb = """
    PREFIX rdae: <http://rdaregistry.info/Elements/m/>
    PREFIX bnf-onto: <http://data.bnf.fr/ontology/bnf-onto/>
    PREFIX rdf: <http://www.w3.org/1999/02/22-rdf-syntax-ns#>
    PREFIX rdfs: <http://www.w3.org/2000/01/rdf-schema#>
    PREFIX foaf: <http://xmlns.com/foaf/0.1/>
    PREFIX rdar: <http://rdvocab.info/RDARelationshipsWEMI/>
    
    SELECT DISTINCT ?source ?propriété ?valeur ?role ?nomFamille ?prénom
    
    WHERE {
        ?source bnf-onto:OCR|rdae:P30016|rdar:electronicReproduction <"""

            # End of the query.
            qe = """> ;  ?propriété ?valeur ;
      		  rdar:expressionManifested ?expression.
    	?expression ?role ?dude .
        ?dude a foaf:Person ;
              foaf:familyName ?nomFamille ;
              foaf:givenName ?prénom.
    }"""
            # Assemble the query with the URL.
            query = qb + l + qe
    
            # Send the query to the SPARQL endpoint and
            # transform the Json results into a Pandas DataFrame.
            results_df = query_db(query, link, "https://data.bnf.fr/sparql")

            # Add the Pandas DataFrame to the list.
            all_dfs.append(results_df)

        # If the line is an ARK URI:
        elif "http://ark.bnf.fr" in l:

            # Transform the URI into the corresponding DataBnF URI.
            link = l.replace("ark.bnf", "data.bnf") + "#about"

            # Define the main parts of the query.
            
            # Beginning of the query.
            qb = """
    PREFIX bnf-onto: <http://data.bnf.fr/ontology/bnf-onto/>
    PREFIX rdf: <http://www.w3.org/1999/02/22-rdf-syntax-ns#>
    PREFIX rdfs: <http://www.w3.org/2000/01/rdf-schema#>
    PREFIX foaf: <http://xmlns.com/foaf/0.1/>
    PREFIX rdar: <http://rdvocab.info/RDARelationshipsWEMI/>
    
    SELECT DISTINCT ?source ?propriété ?valeur ?role ?nomFamille ?prénom
    
    WHERE {
      	BIND ( """

            # End of the query.
            qe = """ as ?source)
        ?source ?propriété ?valeur ;
      		  rdar:expressionManifested ?expression.
    	?expression ?role ?dude .
        ?dude a foaf:Person ;
              foaf:familyName ?nomFamille ;
              foaf:givenName ?prénom.
    }"""
            
            # Assemble the query with the URI.
            query = qb+"<"+link+">"+qe
            
            # Send the query to the SPARQL endpoint and
            # transform the Json results into a Pandas DataFrame.
            results_df = query_db(query, link, "https://data.bnf.fr/sparql")
            # Add the Pandas DataFrame to the list.
            all_dfs.append(results_df)

        # Other options are not yet included.
        else:
            print(f"{l} is not a URL/URI we can use for now.")

    # Assemble all DataFrames into one.
    all_results = pd.concat(all_dfs)

    # Reorder the DataFrame to have one line per book and show it.
    final = pd.DataFrame(reorder(all_results))
    display(final)

    # Write the resulting DataFrame into an XLSX file.
    final.to_excel("iiif_metadata.xlsx")

######################################################

# SEND THE QUERY TO DATA BNF AND RETURN A PANDAS DATAFRAME 

def query_db(query_str, sc, endpoint):

    """
    This function communicates with a SPARQL endpoint
    and returns the response as a Pandas DataFrame.

    :param query_str: A string containing a query written in SPARQL.
    :param sc: A string containing the link around which the query
        was built.
    :param endpoint: A string containing the URL for the SPARQL endpoint.
    
    """

    # Specify the DBPedia endpoint
    sparql = SPARQLWrapper(endpoint)

    # Specify the query.
    sparql.setQuery(query_str)

    # Convert results to JSON format
    sparql.setReturnFormat(JSON)
    result = sparql.query().convert()
    output = result["results"]["bindings"]

    # Return the results as a Pandas DataFrame.
    return to_pd_df(output, sc)

######################################################

# GET THE JSON RESULTS FROM A DATA BNF QUERY AND TURN THEM
# INTO A PANDAS DATAFRAME

def to_pd_df(dicolist, sc):

    """
    This function takes the Json results from a SPARQL query
    and transforms them into a Pandas DataFrame, using keys
    as column heads.

    :param dicolist: The results of a SPARQL query as a list
        of dictionaries [{}]
    :param sc: The link around which the query was built,
        to include in a new column.
    
    """
    
    if len(dicolist) != 0:

        # Make sure you get all the possible column heads.
        keys = []
        for result in dicolist:
            for key in result.keys():
                if key not in keys:
                    keys.append(key)

        # Initiate the dataframe-to-be.
        to_pd_df = {"Source":[]}

        # Initiate each column-to-be.
        for key in keys:
            to_pd_df[key] = []

        # Fill the dataframe-to-be.
        for result in dicolist:
            to_pd_df["Source"].append(sc)
            for key in keys:
                if key in result.keys():
                    to_pd_df[key].append(result[key]["value"])
                else:
                    to_pd_df[key].append(None)
                    
        # Transform the dict in an actual Pandas.DataFrame
        df = pd.DataFrame.from_dict(to_pd_df)
    
        return df
        
    else:
        return None

######################################################

# SORT THE RESULTS THEMATICALLY TO PREPARE THE BIBLIOGRAPHY

def reorder(df):

    """
    This function groups the information to be human-readable
    and remove doubles (DataBnF usually encodes the same information
    several times according to different models).

    :param df: A Pandas DataFrame as produced by the concatenation of
        all DataFrames returned on the links by the the query_db() function.
    
    """

    # Define the columns of the new table.
    # Keys are the new columns.
    # Values are the corresponding DataBnF properties.
    sort = {
        "Title" : ["http://purl.org/dc/terms/title"],
        "Author" :  [
            "http://data.bnf.fr/vocabulary/roles/r70",
            "http://id.loc.gov/vocabulary/relators/aut"],
        "Sc. editor" : [
            "http://data.bnf.fr/vocabulary/roles/r360",
            "http://id.loc.gov/vocabulary/relators/edt"],
        "Contributor" : ["http://purl.org/dc/terms/contributor"],
        "Edition" : [
            "http://rdaregistry.info/Elements/m/#P30133",
            "http://rdvocab.info/Elements/designationOfEdition"],
        "Date" : [
            "http://data.bnf.fr/ontology/bnf-onto/firstYear",
            "http://purl.org/dc/terms/date",
            "http://rdaregistry.info/Elements/m/#P30011",
            "http://rdvocab.info/Elements/dateOfPublicationManifestation"],
        "Place" : [
            "http://rdaregistry.info/Elements/m/#P30279",
            "http://rdvocab.info/Elements/placeOfPublication"],
        "Publisher" : [
            "http://rdaregistry.info/Elements/m/#P30176",
            "http://rdvocab.info/Elements/publishersName"],
        "Publisher (full)" : ["http://purl.org/dc/terms/publisher"],
        "Notes" : [
            "http://rdaregistry.info/Elements/u/#P60470",
            "http://rdvocab.info/Elements/note"],
        "Facsimile" : [
            "http://data.bnf.fr/ontology/bnf-onto/OCR",
            "http://rdaregistry.info/Elements/m/#P30016",
            "http://rdvocab.info/RDARelationshipsWEMI/electronicReproduction"],
        "BnF identifier" : ["http://data.bnf.fr/ontology/bnf-onto/FRBNF"],
        "Description" : ["http://purl.org/dc/terms/description"]
    }

    # Initiate the new DataFrame-to-be.
    new_df = {
        "Title" : [],
        "Author" : [],
        "Sc. editor" : [],
        "Contributor" : [],
        "Other contributor" : [],
        "Edition" : [],
        "Date" : [],
        "Place" : [],
        "Publisher" : [],
        "Publisher (full)" : [],
        "Notes" : [],
        "Source" : [],
        "Facsimile" : [],
        "BnF identifier" : [],
        "Description" : [],
        "Other" : []
    }

    # Loop on each original link.
    for source in np.unique(df["Source"]):

        # Separate the current link information.
        this_one = df[df["Source"] == source]

        # Initiate the corresponding row.
        s = {
        "Title" : [],
        "Author" : [],
        "Sc. editor" : [],
        "Contributor" : [],
        "Other contributor" : [],
        "Edition" : [],
        "Date" : [],
        "Place" : [],
        "Publisher" : [],
        "Publisher (full)" : [],
        "Notes" : [],
        "Source" : [],
        "Facsimile" : [],
        "BnF identifier" : [],
        "Description" : [],
        "Other" : []
        }

        # Loop on all rows of the link's subset.
        for index, row in this_one.iterrows():
            
            # The query basically got out three triplets:
            # ?link ?propriété ?valeur
            # ?correspondingExpression ?role ?dude
            # First we get what ?propriété and ?role are expressed in this row
            # and reconstitute the current contributor's full name.
            p = row["propriété"]
            r = row["role"]
            dude = f"{row['nomFamille']}, {row['prénom']}"

            # All remaining metadata are first sorted into the new columns.
            if p in sort["Title"]:
                s["Title"].append(row["valeur"])
            elif p in sort["Edition"]:
                s["Edition"].append(row["valeur"])
            elif p in sort["Date"]:
                s["Date"].append(row["valeur"].replace("http://data.bnf.fr/date/", "").replace("/",""))
            elif p in sort["Place"]:
                s["Place"].append(row["valeur"])
            elif p in sort["Publisher"]:
                s["Publisher"].append(row["valeur"])
            elif p in sort["Publisher (full)"]:
                s["Publisher (full)"].append(row["valeur"])
            elif p in sort["Notes"]:
                s["Notes"].append(row["valeur"])
            elif p in sort["BnF identifier"]:
                s["BnF identifier"].append(row["valeur"])
            elif p in sort["Facsimile"]:
                s["Facsimile"].append(row["valeur"])
            elif p in sort["Description"]:
                s["Description"].append(row["valeur"])
            else:
                # A column for unforeseen metadata.
                s["Other"].append(f"{row['propriété']} → {row['valeur']}")

            # Now deal with contributors, according to their actual roles.
            if r.strip() in sort["Author"]:
                s["Author"].append(dude)
            elif r in sort["Sc. editor"]:
                s["Sc. editor"].append(dude)
            elif r in sort["Contributor"]:
                s["Contributor"].append(dude)
            else:
                s["Other contributor"].append(f"{r} → {dude}")

        # As an author may also be listed as a contributor,
        # and we want to avoid annoying doubles, compare the lists.
        for c in s["Contributor"]:
            if c in s["Author"] or c in s["Sc. editor"]:
                s["Contributor"].remove(c)

        # Now add all information to the output-to-be.
        for k in s.keys():
            # In case some information is not there.
            if len(s[k]) == 0:
                new_df[k].append(None)
            else:
                # Keep checking on contributor doubles, just in case.
                if k == "Contributor":
                    new_k = []
                    for c in s[k]:
                        if c not in s["Author"] and c not in s["Sc. editor"]:
                            new_k.append(c)
                    if len(new_k) == 0:
                        new_df[k].append(None)
                    else:
                        new_df[k].append(" ; ".join(np.unique(new_k)))
                
                # Otherwise, just add.
                else:
                    new_df[k].append(" ; ".join(np.unique(s[k])))
    return new_df

######################################################

# APPLY OXFORD-STYLE LAYOUT

def author_date(df):

    """
    This function takes an XLSX file as produced by the parse_list() function
    and prepares a DOCX bibliography in Author-Date or Oxford Style layout.

    :param df: A string containing the path to the XLSX input file.
    
    """

    # Initiate the output.
    doc = Document()

    # Acquire and prepare the DataFrame.
    mddd = pd.read_excel("iiif_metadata.xlsx")
    mdd = mddd.replace(np.nan, None)
    md = mdd.sort_values(by="Author")
    
    for idx, row in md.iterrows():

        # Make one paragraph per row and prepare contributor list.
        paragraph = doc.add_paragraph()
        all_dudes = []
        
        if row["Author"] != None:
            authors = row["Author"].split(" ; ")
            for idx, author in enumerate(authors, 1):
                name = author.split(", ")
                all_dudes.append({"order":idx, "fct":"auth", "fam":name[0], "given":name[1]})
                
        if row["Sc. editor"] != None:
            editors = row["Sc. editor"].split(" ; ")
            for idx, editor in enumerate(editors, ):
                name = editor.split(", ")
                all_dudes.append({"order":idx, "fct":"ed", "fam":name[0], "given":name[1]})
                
        if row["Contributor"] != None:
            contribs = row["Contributor"].split(" ; ")
            for idx, contrib in enumerate(contribs, 1):
                name = contrib.split(", ")
                all_dudes.append({"order":idx, "fct":"contrib", "fam":name[0], "given":name[1]})
                
        if row["Other contributor"] != None:
            contribs = row["Other contributor"].split(" ; ")
            for idx, contrib in enumerate(contribs, 1):
                name = contrib.split(", ")
                all_dudes.append({"order":idx, "fct":"contrib", "fam":name[0], "given":name[1]})
                
        # Write the actual contributor list.
        
        if len(all_dudes) == 0:
            paragraph.add_run(text = "Anonyme")
            
        elif len(all_dudes) == 1:
            them = all_dudes[0]
            ln = paragraph.add_run(text = them["fam"])
            ln.font.small_caps = True
            paragraph.add_run(text = ", " + them["given"])
            if them["fct"] == "ed":
                paragraph.add_run(text = "(ed.). ")
                
        elif len(all_dudes) == 2:
            them = all_dudes[0]
            ln = paragraph.add_run(text = them["fam"])
            ln.font.small_caps = True
            paragraph.add_run(text = ", " + them["given"])
            if them["fct"] == "ed":
                paragraph.add_run(text = "(ed.)")
            paragraph.add_run(text = " et ")
            
            them = all_dudes[1]
            ln = paragraph.add_run(text = them["fam"])
            ln.font.small_caps = True
            paragraph.add_run(text = ", " + them["given"])
            if them["fct"] == "ed":
                paragraph.add_run(text = "(ed.)")
                
        elif len(all_dudes) > 2:
            
            print(all_dudes)
            ld = len(all_dudes)
            for idx, dud in enumerate(all_dudes):
    
                if ld-idx >= 3 :
                    ln = paragraph.add_run(text = dud["fam"])
                    ln.font.small_caps = True
                    paragraph.add_run(text = ", " + dud["given"])
                    if dud["fct"] == "ed":
                        paragraph.add_run(text = "(ed.)")
                    paragraph.add_run(text = ", ")
    
                elif ld-idx == 2:
    
                    ln = paragraph.add_run(text = dud["fam"])
                    ln.font.small_caps = True
                    paragraph.add_run(text = ", " + dud["given"])
                    if dud["fct"] == "ed":
                        paragraph.add_run(text = "(ed.)")
                    paragraph.add_run(text = " et ")
                    
                elif ld-idx == 1:
            
                    ln = paragraph.add_run(text = dud["fam"])
                    ln.font.small_caps = True
                    paragraph.add_run(text = ", " + dud["given"])
                    if dud["fct"] == "ed":
                        paragraph.add_run(text = "(ed.)")

        # Add the rest.
        
        paragraph.add_run(text = f". [1e édition ??] ({row['Date']}")
    
        if row["Edition"] != None:
            paragraph.add_run(text = ", " + row["Edition"])
    
        
        paragraph.add_run(text = "). ")
        tit = paragraph.add_run(text = row["Title"])
        tit.italic = True
        paragraph.add_run(text = ". ")
        
        if row["Description"] != None:
            paragraph.add_run(text = row["Description"] + ". ")
        
        if row["Place"] != None:
            pp = row["Place"].split(" ; ")
            places = []
            for p in pp:
                places.append(p.split(" (")[0])
            paragraph.add_run(text = ", ".join(np.unique(places)) + " : ")
    
        if row["Publisher"] != None:
            paragraph.add_run(text = row["Publisher"])
        paragraph.add_run(text = ". ")
    
        if row["Facsimile"] != None:
            paragraph.add_run(text = "En ligne : " + row["Facsimile"] + ".")
    
    # Write the output into a DOCX file.
    doc.save("biblio.docx")

        
























