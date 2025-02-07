# -*- coding: utf-8 -*-

import numpy as np
import pandas as pd
from docx import Document
from SPARQLWrapper import SPARQLWrapper, JSON
from tqdm.notebook import tqdm

######################################################

# WRITE AND SEND THE QUERIES

def parse_list(iiif_list):

    all_dfs = []
    
    with open(iiif_list) as f:
        texte = f.read()
    
    links = texte.split("\n")
    
    for l in tqdm(links):
        if "http://gallica.bnf.fr" in l:
            
            qb = """
    PREFIX rdae: <http://rdaregistry.info/Elements/m/>
    PREFIX bnf-onto: <http://data.bnf.fr/ontology/bnf-onto/>
    PREFIX rdf: <http://www.w3.org/1999/02/22-rdf-syntax-ns#>
    PREFIX rdfs: <http://www.w3.org/2000/01/rdf-schema#>
    PREFIX foaf: <http://xmlns.com/foaf/0.1/>
    PREFIX rdar: <http://rdvocab.info/RDARelationshipsWEMI/>
    
    SELECT DISTINCT ?source ?propriété ?valeur ?role ?nomFamille ?prénom
    
    WHERE {
        ?source bnf-onto:OCR|rdae:P30016|rdar:electronicReproduction <"""
            
            qe = """> ;  ?propriété ?obj1 ;
      		  rdar:expressionManifested ?expression.
    	?expression ?role ?dude .
        ?dude a foaf:Person ;
              foaf:familyName ?nomFamille ;
              foaf:givenName ?prénom.
    }"""
            query = qb + l + qe
    
            results_df = query_db(query, link, "https://data.bnf.fr/sparql")
            
            all_dfs.append(results_df)
            
        elif "http://ark.bnf.fr" in l:
            
            link = l.replace("ark.bnf", "data.bnf") + "#about"
            
            qb = """
    PREFIX bnf-onto: <http://data.bnf.fr/ontology/bnf-onto/>
    PREFIX rdf: <http://www.w3.org/1999/02/22-rdf-syntax-ns#>
    PREFIX rdfs: <http://www.w3.org/2000/01/rdf-schema#>
    PREFIX foaf: <http://xmlns.com/foaf/0.1/>
    PREFIX rdar: <http://rdvocab.info/RDARelationshipsWEMI/>
    
    SELECT DISTINCT ?source ?propriété ?valeur ?role ?nomFamille ?prénom
    
    WHERE {
      	BIND ( """
            
            qe = """ as ?source)
        ?source ?propriété ?valeur ;
      		  rdar:expressionManifested ?expression.
    	?expression ?role ?dude .
        ?dude a foaf:Person ;
              foaf:familyName ?nomFamille ;
              foaf:givenName ?prénom.
    }"""
            
            query = qb+"<"+link+">"+qe
    
            results_df = query_db(query, link, "https://data.bnf.fr/sparql")
            
            all_dfs.append(results_df)
        
        #elif "http://
    
    all_results = pd.concat(all_dfs)
    
    final = pd.DataFrame(reorder(all_results))
    display(final)
    
    final.to_excel("iiif_metadata.xlsx")

######################################################

# SEND THE QUERY TO DATA BNF AND RETURN A PANDAS DATAFRAME 

def query_db(query_str, sc, endpoint):

    # Specify the DBPedia endpoint
    sparql = SPARQLWrapper(endpoint)

    # Query for the description of "Capsaicin", filtered by language
    sparql.setQuery(query_str)

    # Convert results to JSON format
    sparql.setReturnFormat(JSON)
    result = sparql.query().convert()
    output = result["results"]["bindings"]

    return to_pd_df(output, sc)

######################################################

# GET THE JSON RESULTS FROM A DATA BNF QUERY AND TURN THEM
# INTO A PANDAS DATAFRAME

def to_pd_df(dicolist, sc):
    
    if len(dicolist) != 0:

        # Make sure you get all the column heads.
        keys = []
        for result in dicolist:
            for key in result.keys():
                if key not in keys:
                    keys.append(key)

        # Initiate the dataframe-to-be.
        to_pd_df = {"Source":[]}

        for key in keys:
            to_pd_df[key] = []

        # Fill the dataframe-to-be.
        for result in dicolist:
            to_pd_df["Source"].append(sc)
            for key in keys:
                if key in result.keys():
                    to_pd_df[key].append(result[key]["value"])
                else:
                    to_pd_df[key].append(None)
                    
        # Transform the dict in an actual Pandas.DataFrame
        df = pd.DataFrame.from_dict(to_pd_df)
    
        return df
    else:
        return None

######################################################

# SORT THE RESULTS THEMATICALLY TO PREPARE THE BIBLIOGRAPHY

def reorder(df):

    sort = {
        "Title" : ["http://purl.org/dc/terms/title"],
        "Author" :  [
            "http://data.bnf.fr/vocabulary/roles/r70",
            "http://id.loc.gov/vocabulary/relators/aut"],
        "Sc. editor" : [
            "http://data.bnf.fr/vocabulary/roles/r360",
            "http://id.loc.gov/vocabulary/relators/edt"],
        "Contributor" : ["http://purl.org/dc/terms/contributor"],
        "Edition" : [
            "http://rdaregistry.info/Elements/m/#P30133",
            "http://rdvocab.info/Elements/designationOfEdition"],
        "Date" : [
            "http://data.bnf.fr/ontology/bnf-onto/firstYear",
            "http://purl.org/dc/terms/date",
            "http://rdaregistry.info/Elements/m/#P30011",
            "http://rdvocab.info/Elements/dateOfPublicationManifestation"],
        "Place" : [
            "http://rdaregistry.info/Elements/m/#P30279",
            "http://rdvocab.info/Elements/placeOfPublication"],
        "Publisher" : [
            "http://rdaregistry.info/Elements/m/#P30176",
            "http://rdvocab.info/Elements/publishersName"],
        "Publisher (full)" : ["http://purl.org/dc/terms/publisher"],
        "Notes" : [
            "http://rdaregistry.info/Elements/u/#P60470",
            "http://rdvocab.info/Elements/note"],
        "Facsimile" : [
            "http://data.bnf.fr/ontology/bnf-onto/OCR",
            "http://rdaregistry.info/Elements/m/#P30016",
            "http://rdvocab.info/RDARelationshipsWEMI/electronicReproduction"],
        "BnF identifier" : ["http://data.bnf.fr/ontology/bnf-onto/FRBNF"],
        "Description" : ["http://purl.org/dc/terms/description"]
    }

    new_df = {
        "Title" : [],
        "Author" : [],
        "Sc. editor" : [],
        "Contributor" : [],
        "Other contributor" : [],
        "Edition" : [],
        "Date" : [],
        "Place" : [],
        "Publisher" : [],
        "Publisher (full)" : [],
        "Notes" : [],
        "Source" : [],
        "Facsimile" : [],
        "BnF identifier" : [],
        "Description" : [],
        "Other" : []
    }

    for source in np.unique(df["Source"]):
        
        this_one = df[df["Source"] == source]
        s = {
        "Title" : [],
        "Author" : [],
        "Sc. editor" : [],
        "Contributor" : [],
        "Other contributor" : [],
        "Edition" : [],
        "Date" : [],
        "Place" : [],
        "Publisher" : [],
        "Publisher (full)" : [],
        "Notes" : [],
        "Source" : [],
        "Facsimile" : [],
        "BnF identifier" : [],
        "Description" : [],
        "Other" : []
        }
        
        for index, row in this_one.iterrows():

            p = row["propriété"]
            r = row["role"]
            dude = f"{row['nomFamille']}, {row['prénom']}"
            
            if p in sort["Title"]:
                s["Title"].append(row["valeur"])
            elif p in sort["Edition"]:
                s["Edition"].append(row["valeur"])
            elif p in sort["Date"]:
                s["Date"].append(row["valeur"].replace("http://data.bnf.fr/date/", "").replace("/",""))
            elif p in sort["Place"]:
                s["Place"].append(row["valeur"])
            elif p in sort["Publisher"]:
                s["Publisher"].append(row["valeur"])
            elif p in sort["Publisher (full)"]:
                s["Publisher (full)"].append(row["valeur"])
            elif p in sort["Notes"]:
                s["Notes"].append(row["valeur"])
            elif p in sort["BnF identifier"]:
                s["BnF identifier"].append(row["valeur"])
            elif p in sort["Facsimile"]:
                s["Facsimile"].append(row["valeur"])
            elif p in sort["Description"]:
                s["Description"].append(row["valeur"])
            else:
                s["Other"].append(f"{row['propriété']} → {row['valeur']}")

            if r.strip() in sort["Author"]:
                s["Author"].append(dude)
            elif r in sort["Sc. editor"]:
                s["Sc. editor"].append(dude)
            elif r in sort["Contributor"]:
                s["Contributor"].append(dude)
            else:
                s["Other contributor"].append(f"{r} → {dude}")
                
        for c in s["Contributor"]:
            if c in s["Author"] or c in s["Sc. editor"]:
                s["Contributor"].remove(c)
        
        for k in s.keys():
            if len(s[k]) == 0:
                new_df[k].append(None)
            else:
                if k == "Contributor":
                    new_k = []
                    for c in s[k]:
                        if c not in s["Author"] and c not in s["Sc. editor"]:
                            new_k.append(c)
                    if len(new_k) == 0:
                        new_df[k].append(None)
                    else:
                        new_df[k].append(" ; ".join(np.unique(new_k)))
                else:
                    new_df[k].append(" ; ".join(np.unique(s[k])))
    return new_df

######################################################

# APPLY LAYOUT

def layout(df):

    doc = Document()

    mddd = pd.read_excel("iiif_metadata.xlsx")
    mdd = mddd.replace(np.nan, None)
    md = mdd.sort_values(by="Author")
    
    for idx, row in md.iterrows():
    
        paragraph = doc.add_paragraph()
        all_dudes = []
        
        if row["Author"] != None:
            authors = row["Author"].split(" ; ")
            for idx, author in enumerate(authors, 1):
                name = author.split(", ")
                all_dudes.append({"order":idx, "fct":"auth", "fam":name[0], "given":name[1]})
                
        if row["Sc. editor"] != None:
            editors = row["Sc. editor"].split(" ; ")
            for idx, editor in enumerate(editors, ):
                name = editor.split(", ")
                all_dudes.append({"order":idx, "fct":"ed", "fam":name[0], "given":name[1]})
                
        if row["Contributor"] != None:
            contribs = row["Contributor"].split(" ; ")
            for idx, contrib in enumerate(contribs, 1):
                name = contrib.split(", ")
                all_dudes.append({"order":idx, "fct":"contrib", "fam":name[0], "given":name[1]})
                
        if row["Other contributor"] != None:
            contribs = row["Other contributor"].split(" ; ")
            for idx, contrib in enumerate(contribs, 1):
                name = contrib.split(", ")
                all_dudes.append({"order":idx, "fct":"contrib", "fam":name[0], "given":name[1]})
                
    
        if len(all_dudes) == 0:
            paragraph.add_run(text = "Anonyme")
            
        elif len(all_dudes) == 1:
            them = all_dudes[0]
            ln = paragraph.add_run(text = them["fam"])
            ln.font.small_caps = True
            paragraph.add_run(text = ", " + them["given"])
            if them["fct"] == "ed":
                paragraph.add_run(text = "(ed.). ")
                
        elif len(all_dudes) == 2:
            them = all_dudes[0]
            ln = paragraph.add_run(text = them["fam"])
            ln.font.small_caps = True
            paragraph.add_run(text = ", " + them["given"])
            if them["fct"] == "ed":
                paragraph.add_run(text = "(ed.)")
            paragraph.add_run(text = " et ")
            
            them = all_dudes[1]
            ln = paragraph.add_run(text = them["fam"])
            ln.font.small_caps = True
            paragraph.add_run(text = ", " + them["given"])
            if them["fct"] == "ed":
                paragraph.add_run(text = "(ed.)")
                
        elif len(all_dudes) > 2:
            
            print(all_dudes)
            ld = len(all_dudes)
            for idx, dud in enumerate(all_dudes):
    
                if ld-idx >= 3 :
                    ln = paragraph.add_run(text = dud["fam"])
                    ln.font.small_caps = True
                    paragraph.add_run(text = ", " + dud["given"])
                    if dud["fct"] == "ed":
                        paragraph.add_run(text = "(ed.)")
                    paragraph.add_run(text = ", ")
    
                elif ld-idx == 2:
    
                    ln = paragraph.add_run(text = dud["fam"])
                    ln.font.small_caps = True
                    paragraph.add_run(text = ", " + dud["given"])
                    if dud["fct"] == "ed":
                        paragraph.add_run(text = "(ed.)")
                    paragraph.add_run(text = " et ")
                    
                elif ld-idx == 1:
            
                    ln = paragraph.add_run(text = dud["fam"])
                    ln.font.small_caps = True
                    paragraph.add_run(text = ", " + dud["given"])
                    if dud["fct"] == "ed":
                        paragraph.add_run(text = "(ed.)")
    
        paragraph.add_run(text = f". [1e édition ??] ({row['Date']}")
    
        if row["Edition"] != None:
            paragraph.add_run(text = ", " + row["Edition"])
    
        
        paragraph.add_run(text = "). ")
        tit = paragraph.add_run(text = row["Title"])
        tit.italic = True
        paragraph.add_run(text = ". ")
        
        if row["Description"] != None:
            paragraph.add_run(text = row["Description"] + ". ")
        
        if row["Place"] != None:
            pp = row["Place"].split(" ; ")
            places = []
            for p in pp:
                places.append(p.split(" (")[0])
            paragraph.add_run(text = ", ".join(np.unique(places)) + " : ")
    
        if row["Publisher"] != None:
            paragraph.add_run(text = row["Publisher"])
        paragraph.add_run(text = ". ")
    
        if row["Facsimile"] != None:
            paragraph.add_run(text = "En ligne : " + row["Facsimile"] + ".")
    
    
    doc.save("biblio.docx")

        
























